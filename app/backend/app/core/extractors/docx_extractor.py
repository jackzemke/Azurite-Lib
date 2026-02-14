"""
DOCX (Word Document) Extractor.

Extracts text and metadata from Microsoft Word documents.
Preserves document order (paragraphs and tables interleaved correctly),
converts tables to Markdown, and detects heading styles.
"""

from docx import Document
from docx.oxml.ns import qn
from pathlib import Path
from typing import Dict, List
import logging

logger = logging.getLogger(__name__)


def _docx_table_to_markdown(table) -> str:
    """Convert a python-docx Table object to Markdown format.

    Args:
        table: A python-docx Table object

    Returns:
        Markdown table string, or empty string if table is empty.
    """
    rows = []
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        rows.append(cells)

    if not rows or len(rows) < 2:
        # Single-row table or empty — just return as text
        if rows:
            return " | ".join(rows[0])
        return ""

    # Skip tables where all cells are empty
    if all(all(c == "" for c in row) for row in rows):
        return ""

    # Determine column count from widest row
    n_cols = max(len(row) for row in rows)

    # Pad rows to uniform width
    for row in rows:
        while len(row) < n_cols:
            row.append("")

    # Build Markdown table
    header = rows[0]
    separator = ["---"] * n_cols
    body = rows[1:]

    lines = []
    lines.append("| " + " | ".join(header) + " |")
    lines.append("| " + " | ".join(separator) + " |")
    for row in body:
        lines.append("| " + " | ".join(row) + " |")

    return "\n".join(lines)


def _get_heading_level(paragraph) -> int:
    """Get the heading level of a paragraph (0 = not a heading).

    Args:
        paragraph: A python-docx Paragraph object

    Returns:
        Heading level (1-9) or 0 if not a heading.
    """
    style_name = paragraph.style.name if paragraph.style else ""
    if style_name.startswith("Heading"):
        try:
            return int(style_name.split()[-1])
        except (ValueError, IndexError):
            return 1
    elif style_name == "Title":
        return 1
    elif style_name == "Subtitle":
        return 2
    return 0


class DOCXExtractor:
    """Extract text and metadata from DOCX files."""

    def extract(self, file_path: Path) -> Dict:
        """
        Extract text from DOCX file, preserving document order.

        Iterates through the document body element to maintain the correct
        interleaving of paragraphs and tables. Tables are converted to
        Markdown format. Heading styles are converted to Markdown headings.

        Args:
            file_path: Path to DOCX file

        Returns:
            Dict with:
                - pages: List of dicts (simulated pages, one per major section)
                - total_pages: int (number of sections)
                - requires_ocr: bool (always False for DOCX)
                - metadata: dict (core properties)
        """
        try:
            doc = Document(file_path)

            # Extract metadata
            core_props = doc.core_properties
            metadata = {
                "title": core_props.title or "",
                "author": core_props.author or "",
                "created": str(core_props.created) if core_props.created else None,
                "modified": str(core_props.modified) if core_props.modified else None,
            }

            # Build a lookup of table elements for quick matching
            table_elements = {}
            for table in doc.tables:
                table_elements[table._element] = table

            # Iterate body elements in document order
            content_blocks = []
            for child in doc.element.body:
                if child.tag == qn('w:p'):
                    # Paragraph element — find the matching Paragraph object
                    for para in doc.paragraphs:
                        if para._element is child:
                            text = para.text.strip()
                            if text:
                                heading_level = _get_heading_level(para)
                                if heading_level > 0:
                                    content_blocks.append(f"{'#' * heading_level} {text}")
                                else:
                                    content_blocks.append(text)
                            break

                elif child.tag == qn('w:tbl'):
                    # Table element
                    table = table_elements.get(child)
                    if table:
                        md = _docx_table_to_markdown(table)
                        if md:
                            content_blocks.append(md)

            # Simulate pages: group content blocks
            page_size = 10
            pages = []
            for i in range(0, len(content_blocks), page_size):
                page_blocks = content_blocks[i:i + page_size]
                page_text = "\n\n".join(page_blocks)
                pages.append({
                    "page_num": (i // page_size) + 1,
                    "text": page_text,
                    "bbox_available": False,
                    "width": None,
                    "height": None,
                })

            # If no content, create one empty page
            if not pages:
                pages = [{
                    "page_num": 1,
                    "text": "",
                    "bbox_available": False,
                    "width": None,
                    "height": None,
                }]

            return {
                "pages": pages,
                "total_pages": len(pages),
                "requires_ocr": False,
                "metadata": metadata,
            }

        except Exception as e:
            logger.error(f"Failed to extract DOCX {file_path}: {e}")
            raise
